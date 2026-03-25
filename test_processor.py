"""
Smoke tests – creates a .docx with comments, highlights, paragraph shading
and images, runs the processor, and checks the output.
"""

import os
import tempfile

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from processor import remove_comments, remove_highlighting, remove_images_after_page_one


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _add_run_highlight(run, color='yellow'):
    rpr = run._r.get_or_add_rPr()
    hl = etree.SubElement(rpr, qn('w:highlight'))
    hl.set(qn('w:val'), color)


def _add_paragraph_shading(paragraph, fill='FFFF00'):
    """Add paragraph-level shading (whole-paragraph background colour)."""
    ppr = paragraph._p.get_or_add_pPr()
    shd = etree.SubElement(ppr, qn('w:shd'))
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), fill)


def _add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def _count_drawings(doc):
    return len(doc.element.body.findall(f'.//{qn("w:drawing")}'))


def _count_highlights(doc):
    return len(doc.element.body.findall(f'.//{qn("w:highlight")}'))


def _count_paragraph_shadings(doc):
    count = 0
    for p in doc.element.body.iter(qn('w:p')):
        ppr = p.find(qn('w:pPr'))
        if ppr is not None and ppr.find(qn('w:shd')) is not None:
            count += 1
    return count


def _count_comment_markers(doc):
    count = 0
    for tag in ('w:commentRangeStart', 'w:commentRangeEnd', 'w:commentReference'):
        count += len(doc.element.body.findall(f'.//{qn(tag)}'))
    return count


# ---------------------------------------------------------------------------
# Test document
# ---------------------------------------------------------------------------

def build_test_doc(path):
    doc = Document()

    # Page 1 ─────────────────────────────────────────────────────────────────
    p1 = doc.add_paragraph('Pagina 1 tekst')
    run1 = p1.add_run(' gearceerd run')
    _add_run_highlight(run1, 'yellow')

    p1_shaded = doc.add_paragraph('Gearceerde alinea pagina 1')
    _add_paragraph_shading(p1_shaded, 'FFFF00')

    # Comment marker
    crs = OxmlElement('w:commentRangeStart')
    crs.set(qn('w:id'), '1')
    p1._p.insert(0, crs)

    # Inline image placeholder on page 1
    p_img1 = doc.add_paragraph('Afbeelding pagina 1:')
    drawing = OxmlElement('w:drawing')
    p_img1._p.append(drawing)

    _add_page_break(doc)

    # Page 2+ ─────────────────────────────────────────────────────────────────
    p2 = doc.add_paragraph('Pagina 2 tekst')

    # Image that MUST be removed
    drawing2 = OxmlElement('w:drawing')
    p2._p.append(drawing2)

    run2 = p2.add_run(' ook gearceerd')
    _add_run_highlight(run2, 'green')

    p2_shaded = doc.add_paragraph('Gearceerde alinea pagina 2')
    _add_paragraph_shading(p2_shaded, '00FF00')

    doc.save(path)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

def test_all():
    with tempfile.TemporaryDirectory() as tmpdir:
        path = os.path.join(tmpdir, 'test.docx')
        build_test_doc(path)

        doc = Document(path)

        # Verify preconditions
        assert _count_highlights(doc) == 2, f"Setup: expected 2 highlights, got {_count_highlights(doc)}"
        assert _count_paragraph_shadings(doc) == 2, f"Setup: expected 2 para shadings, got {_count_paragraph_shadings(doc)}"
        assert _count_drawings(doc) == 2, f"Setup: expected 2 drawings, got {_count_drawings(doc)}"
        assert _count_comment_markers(doc) >= 1, "Setup: expected ≥1 comment marker"

        # Apply transforms
        remove_comments(doc)
        remove_highlighting(doc)
        remove_images_after_page_one(doc)

        hl = _count_highlights(doc)
        assert hl == 0, f"Run highlights not removed: {hl} remain"
        print(f"  [OK] Run highlights removed: 0 remain")

        ps = _count_paragraph_shadings(doc)
        assert ps == 0, f"Paragraph shadings not removed: {ps} remain"
        print(f"  [OK] Paragraph shadings removed: 0 remain")

        dr = _count_drawings(doc)
        assert dr == 1, f"Expected 1 drawing (page 1 only), got {dr}"
        print(f"  [OK] Drawings after page 1 removed: 1 remains (page 1)")

        cm = _count_comment_markers(doc)
        assert cm == 0, f"Comment markers not removed: {cm} remain"
        print(f"  [OK] Comment markers removed: 0 remain")

    print("\nAll tests passed.")


if __name__ == '__main__':
    test_all()
