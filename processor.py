"""
Word-to-PDF processor.

Transforms a .docx file into a print-ready PDF by:
  1. Removing all comments
  2. Removing all highlighting (arceringen) while preserving text colour
  3. Removing all images except those on page 1
  4. Keeping page-number footer intact
"""

import logging
import os
import shutil
import subprocess
import sys
import tempfile
import threading
import urllib.parse
import uuid

from docx import Document
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Headless bootstrap — auto-install missing system packages on Linux
# ---------------------------------------------------------------------------

log = logging.getLogger(__name__)

_bootstrap_lock = threading.Lock()
_bootstrap_done = False


def _try_apt_install(*packages: str) -> bool:
    """
    Attempt to install one or more Debian/Ubuntu packages via apt-get.
    Runs 'apt-get update' first so the package can be found even in fresh
    containers where the apt cache has never been populated.
    Returns True if the install command succeeded.
    Logs failures but never raises so callers stay unaffected.
    """
    if sys.platform != 'linux' or not shutil.which('apt-get'):
        return False
    apt_env = {**os.environ, 'DEBIAN_FRONTEND': 'noninteractive'}
    try:
        # Refresh index first — a cold apt cache causes "Unable to locate package"
        subprocess.run(
            ['apt-get', 'update', '-q'],
            capture_output=True, timeout=120, env=apt_env,
        )
        r = subprocess.run(
            ['apt-get', 'install', '-y', '--no-install-recommends', *packages],
            capture_output=True,
            timeout=180,
            env=apt_env,
        )
        if r.returncode != 0:
            log.warning(
                'apt-get install %s failed (rc=%d): %s',
                ' '.join(packages), r.returncode,
                r.stderr.decode(errors='replace') if isinstance(r.stderr, bytes) else r.stderr,
            )
        return r.returncode == 0
    except PermissionError:
        log.warning(
            'apt-get install %s skipped: no root privileges. '
            'Run manually: sudo apt-get install %s',
            ' '.join(packages), ' '.join(packages),
        )
        return False
    except Exception as exc:
        log.warning('apt-get install %s failed: %s', ' '.join(packages), exc)
        return False


def _lo_works_headless() -> bool:
    """
    Return True if LibreOffice can actually convert a document without a display.

    We intentionally do NOT run 'libreoffice --version' here: that command
    skips VCL initialisation entirely and returns 0 even when the headless
    plugin is missing — giving a false-positive that tricks bootstrap into
    thinking everything is fine.

    Instead we check directly whether the svp VCL plugin library exists on
    disk (installed by the libreoffice-headless package) or whether xvfb-run
    is available as a fallback display server.

    We search all known install locations:
      - /usr/lib/libreoffice     (Debian/Ubuntu standard)
      - /usr/lib64/libreoffice   (RHEL/Fedora/openSUSE)
      - /opt/libreoffice*        (manually downloaded tarball)
      - /snap/libreoffice/*      (Snap package)
      - directory next to the libreoffice binary (custom/portable installs)
    """
    if sys.platform != 'linux':
        return True  # macOS uses its native renderer; no special setup needed

    import glob as _glob

    # Candidate glob patterns — order is fastest-match-first
    patterns = [
        '/usr/lib/libreoffice/program/libvclplug_svp*.so',
        '/usr/lib64/libreoffice/program/libvclplug_svp*.so',
        '/opt/libreoffice*/program/libvclplug_svp*.so',
        '/snap/libreoffice/*/usr/lib/libreoffice/program/libvclplug_svp*.so',
    ]
    if any(_glob.glob(p) for p in patterns):
        return True

    # Also check relative to the actual soffice binary (handles custom/portable installs)
    lo_bin = shutil.which('libreoffice') or shutil.which('soffice')
    if lo_bin:
        prog_dir = os.path.dirname(os.path.realpath(lo_bin))
        if _glob.glob(os.path.join(prog_dir, 'libvclplug_svp*.so')):
            return True

    # xvfb-run provides a virtual X11 display as a fallback
    return bool(shutil.which('xvfb-run'))


def bootstrap_headless_libreoffice() -> None:
    """
    Ensure LibreOffice can run headlessly. Safe to call from multiple threads
    or processes: the inner work runs exactly once per process thanks to a
    module-level lock + flag; apt-get's own advisory lock prevents concurrent
    installs across processes.

    On Linux: verifies that the svp VCL plugin is present (or xvfb-run is
    available) and installs the missing package automatically via apt-get.

    Silent on macOS (native renderer; no extra packages needed) and when
    apt-get is unavailable (non-Debian systems).

    The first cold start may take 30–60 s while the package downloads;
    every subsequent start is instant because the package stays installed.
    """
    global _bootstrap_done

    # Fast path — already done in this process (no lock needed for read)
    if _bootstrap_done:
        return

    if sys.platform != 'linux':
        _bootstrap_done = True
        return

    with _bootstrap_lock:
        # Re-check inside the lock (double-checked locking)
        if _bootstrap_done:
            return

        try:
            if _lo_works_headless():
                return  # Already fine — nothing to do

            log.warning(
                'LibreOffice headless check failed. '
                'Attempting to install libreoffice-headless automatically…'
            )

            if _try_apt_install('libreoffice-headless'):
                if _lo_works_headless():
                    log.info('libreoffice-headless installed successfully.')
                    return
                log.warning('libreoffice-headless installed but check still fails. Trying xvfb…')

            if _try_apt_install('xvfb'):
                log.info('xvfb installed as fallback display backend.')
                return

            log.error(
                'Could not fix LibreOffice display issue automatically. '
                'Run manually: sudo apt-get install libreoffice-headless'
            )
        finally:
            # Mark done regardless of outcome so we never re-run in this process
            _bootstrap_done = True


# ---------------------------------------------------------------------------
# LibreOffice binary detection
# ---------------------------------------------------------------------------

_LIBREOFFICE_BINARY: str | None = None  # module-level cache


def _ensure_executable(path: str) -> None:
    """Ensure a binary has the executable bit set (may be lost via shutil.copy2)."""
    import stat as _stat
    try:
        mode = os.stat(path).st_mode
        if not (mode & _stat.S_IXUSR):
            os.chmod(path, mode | _stat.S_IXUSR | _stat.S_IXGRP | _stat.S_IXOTH)
    except OSError:
        pass  # Best-effort; failure will surface as a subprocess exec error


def _find_libreoffice() -> str:
    """
    Return the path to the LibreOffice soffice binary.

    Priority on macOS (standalone .app build):
      1. User-library version  – updated automatically by updater.py
      2. Bundled version       – shipped inside the .app at build time
      3. System-installed      – /Applications/LibreOffice.app or PATH

    On Linux the function expects 'libreoffice' on PATH.

    The result is cached at module level so repeated conversions skip
    the filesystem probing overhead.
    """
    global _LIBREOFFICE_BINARY
    if _LIBREOFFICE_BINARY and os.path.isfile(_LIBREOFFICE_BINARY):
        return _LIBREOFFICE_BINARY

    import shutil as _shutil

    if sys.platform == 'darwin':
        # 1 + 2: ask the updater (handles both user-library and bundled paths)
        try:
            from updater import get_active_soffice
            path = get_active_soffice()
            if path and os.path.isfile(path):
                _ensure_executable(path)
                _LIBREOFFICE_BINARY = path
                return path
        except ImportError:
            pass  # updater not available (e.g. running tests on Linux)

        # 3: system-installed fallback
        candidates = [
            _shutil.which('libreoffice'),
            _shutil.which('soffice'),
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '/opt/homebrew/bin/libreoffice',
            '/usr/local/bin/libreoffice',
        ]
        for path in candidates:
            if path and os.path.isfile(path):
                _ensure_executable(path)
                _LIBREOFFICE_BINARY = path
                return path

        raise RuntimeError(
            'LibreOffice niet gevonden.\n'
            'Download en installeer het via https://www.libreoffice.org/download/'
        )

    # Linux / other: expect 'libreoffice' on PATH
    binary = _shutil.which('libreoffice') or _shutil.which('soffice')
    if not binary:
        raise RuntimeError(
            'LibreOffice niet gevonden. Installeer het via:\n'
            '  sudo apt-get install libreoffice-writer'
        )
    _LIBREOFFICE_BINARY = binary
    return binary


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

# Markup Compatibility namespace (Word 2010+ wraps images in mc:AlternateContent)
_MC_NS = 'http://schemas.openxmlformats.org/markup-compatibility/2006'
_MC_ALT = f'{{{_MC_NS}}}AlternateContent'


def _find_first_pagebreak_index(doc):
    """
    Return the index (among body children) of the element that contains the
    first page boundary.

    A page boundary is detected by:
      1. An explicit <w:br w:type="page"/> inside a run, OR
      2. A <w:sectPr> inside a paragraph's <w:pPr> with type nextPage /
         evenPage / oddPage (or no type, which defaults to nextPage).

    The body-level <w:sectPr> (always the sole direct-child sectPr of <w:body>)
    describes the document's last section and is intentionally skipped.

    Returns None if no page break is found (treat whole document as page 1).
    """
    body = doc.element.body
    children = list(body)
    # The body-level sectPr is a direct child; find() only searches direct children.
    body_sectPr = body.find(qn('w:sectPr'))

    for idx, child in enumerate(children):
        # Explicit page break
        for br in child.iter(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                return idx

        # Section break inside a paragraph (pPr/sectPr)
        for sectPr in child.iter(qn('w:sectPr')):
            if sectPr is body_sectPr:
                continue
            t = sectPr.find(qn('w:type'))
            # No <w:type> means nextPage (Word default)
            if t is None or t.get(qn('w:val')) in ('nextPage', 'evenPage', 'oddPage'):
                return idx

    return None


# ---------------------------------------------------------------------------
# 1. Remove comments
# ---------------------------------------------------------------------------

def remove_comments(doc):
    """
    Strip all comment markup from the document body and remove the comments
    relationship part so the converted PDF has no comment side-panels.
    """
    body = doc.element.body

    for tag in (qn('w:commentRangeStart'), qn('w:commentRangeEnd')):
        for elem in body.findall(f'.//{tag}'):
            parent = elem.getparent()
            if parent is not None:
                parent.remove(elem)

    # Runs whose sole content is a commentReference can be dropped entirely
    for run in body.findall(f'.//{qn("w:r")}'):
        non_rpr = [c for c in run if c.tag != qn('w:rPr')]
        if len(non_rpr) == 1 and non_rpr[0].tag == qn('w:commentReference'):
            parent = run.getparent()
            if parent is not None:
                parent.remove(run)

    # Drop the comments part from the package so LibreOffice shows no panel
    try:
        rels_to_drop = [
            key for key, rel in doc.part.rels.items()
            if 'comment' in rel.reltype.lower()
        ]
        for key in rels_to_drop:
            del doc.part.rels[key]
    except Exception:
        pass


# ---------------------------------------------------------------------------
# 2. Remove highlighting
# ---------------------------------------------------------------------------

def remove_highlighting(doc):
    """
    Remove all highlighting/shading from every run and paragraph, leaving
    <w:color> (text colour) and all other properties intact.
    """
    # Body paragraphs
    for paragraph in doc.paragraphs:
        _strip_highlight_from_paragraph(paragraph)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _strip_highlight_from_paragraph(paragraph)

    # Headers and footers
    for section in doc.sections:
        for hf in (
            section.header, section.footer,
            section.first_page_header, section.first_page_footer,
            section.even_page_header, section.even_page_footer,
        ):
            try:
                for paragraph in hf.paragraphs:
                    _strip_highlight_from_paragraph(paragraph)
            except Exception:
                pass


def _strip_highlight_from_paragraph(paragraph):
    # Run-level highlight and shading
    for run in paragraph.runs:
        rpr = run._r.find(qn('w:rPr'))
        if rpr is not None:
            for hl in rpr.findall(qn('w:highlight')):
                rpr.remove(hl)
            for shd in rpr.findall(qn('w:shd')):
                if shd.get(qn('w:val'), 'clear') in ('clear', 'solid'):
                    rpr.remove(shd)

    # Paragraph-level shading (e.g. whole-paragraph background highlight)
    ppr = paragraph._p.find(qn('w:pPr'))
    if ppr is not None:
        for shd in ppr.findall(qn('w:shd')):
            if shd.get(qn('w:val'), 'clear') in ('clear', 'solid'):
                ppr.remove(shd)


# ---------------------------------------------------------------------------
# 3. Remove images after page 1
# ---------------------------------------------------------------------------

def remove_images_after_page_one(doc):
    """
    Remove all images that appear after the first page break.
    Images on page 1 are preserved.
    """
    body = doc.element.body
    children = list(body)
    pagebreak_idx = _find_first_pagebreak_index(doc)

    if pagebreak_idx is None:
        return  # No page break → whole doc is page 1, nothing to remove

    for child in children[pagebreak_idx + 1:]:
        _remove_drawing_elements(child)


# Image container tags — includes mc:AlternateContent (Word 2010+ images)
_DRAWING_TAGS = frozenset([
    qn('w:drawing'),
    qn('w:pict'),
    _MC_ALT,
    '{urn:schemas-microsoft-com:vml}shape',
    '{urn:schemas-microsoft-com:vml}image',
])


def _remove_drawing_elements(element):
    """Remove all image/drawing containers from an XML subtree."""
    for tag in _DRAWING_TAGS:
        for node in element.findall(f'.//{tag}'):
            parent = node.getparent()
            if parent is not None:
                parent.remove(node)


# ---------------------------------------------------------------------------
# 4. Convert processed .docx → PDF via LibreOffice
# ---------------------------------------------------------------------------

def _lo_cmd(binary: str, profile_dir: str, output_dir: str, docx_path: str) -> list:
    return [
        binary,
        '--headless',
        '--norestore',
        '--nofirststartwizard',
        f'-env:UserInstallation=file://{urllib.parse.quote(profile_dir)}',
        '--convert-to', 'pdf',
        '--outdir', output_dir,
        docx_path,
    ]


def _lo_env(extra: dict | None = None) -> dict:
    """
    Build a clean environment for the LibreOffice subprocess.

    The critical problem on macOS inside a PyInstaller-frozen app:
    PyInstaller injects DYLD_LIBRARY_PATH (and related DYLD_* vars) pointing
    to its _MEIPASS extraction directory so the frozen Python host finds its
    own bundled dylibs.  LibreOffice inherits this environment as a subprocess.
    When LibreOffice calls dlopen() to load its VCL plugin (libvclplug_svp.dylib)
    the dynamic linker starts its search in _MEIPASS — where it finds
    incompatible versions of system frameworks — and the load fails with
    "no suitable windowing system found, exiting".

    Fix: strip every DYLD_* variable that PyInstaller may have set before
    handing the environment to LibreOffice.  This is safe because LibreOffice
    uses @rpath / @executable_path (baked into the binary at link time) to
    find its own frameworks, so it does not need DYLD_LIBRARY_PATH.
    """
    env = os.environ.copy()
    for var in (
        'DYLD_LIBRARY_PATH',
        'DYLD_FRAMEWORK_PATH',
        'DYLD_FALLBACK_LIBRARY_PATH',
        'DYLD_INSERT_LIBRARIES',
        # PyInstaller also sets these; they should not reach LibreOffice
        'PYTHONPATH',
        'PYTHONHOME',
    ):
        env.pop(var, None)
    if extra:
        env.update(extra)
    return env


def _is_display_error(stderr: str) -> bool:
    markers = ('windowing system', 'cannot connect to x', 'no display')
    low = stderr.lower()
    return any(m in low for m in markers)


def _validate_pdf(pdf_path: str, source: str = '') -> None:
    """Raise RuntimeError if pdf_path is missing, too small, or has no %PDF header."""
    if not os.path.exists(pdf_path):
        raise RuntimeError(f'{source}Geen PDF gegenereerd op het verwachte pad.')
    size = os.path.getsize(pdf_path)
    if size < 1024:
        raise RuntimeError(f'{source}Ongeldig PDF-bestand ({size} bytes).')
    with open(pdf_path, 'rb') as f:
        if f.read(4) != b'%PDF':
            raise RuntimeError(f'{source}Beschadigd PDF-bestand (ongeldige header).')


def _convert_with_weasyprint(docx_path: str, output_dir: str) -> str:
    """
    Convert DOCX → PDF using mammoth (DOCX→HTML) + WeasyPrint (HTML→PDF).

    Completely pure Python — no external processes, no display, no VCL plugins.
    Works headlessly on macOS (PyInstaller frozen app) and Linux (Docker).
    Requires: pip install mammoth weasyprint
    """
    import base64 as _b64

    import mammoth          # type: ignore
    import weasyprint       # type: ignore

    base     = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(output_dir, base + '.pdf')

    # ── Extract header / footer text via python-docx ─────────────────────────
    header_html = footer_html = ''
    try:
        _doc = Document(docx_path)
        for _sec in _doc.sections:
            if _sec.header and not _sec.header.is_linked_to_previous:
                _ht = '\n'.join(p.text for p in _sec.header.paragraphs if p.text.strip())
                if _ht:
                    header_html = f'<div class="doc-header">{_ht}</div>'
            if _sec.footer and not _sec.footer.is_linked_to_previous:
                _ft = '\n'.join(p.text for p in _sec.footer.paragraphs if p.text.strip())
                if _ft:
                    footer_html = f'<div class="doc-footer">{_ft}</div>'
            break  # first section only
    except Exception:
        pass  # best-effort; missing header/footer is acceptable

    # ── DOCX body → HTML via mammoth ─────────────────────────────────────────
    def _embed_image(image):
        with image.open() as f:
            data = _b64.b64encode(f.read()).decode()
        return {'src': f'data:{image.content_type};base64,{data}'}

    with open(docx_path, 'rb') as f:
        _result = mammoth.convert_to_html(
            f,
            convert_image=mammoth.images.img_element(_embed_image),
        )
    body_html = _result.value

    # ── Minimal A4 stylesheet ─────────────────────────────────────────────────
    css = """
        @page { size: A4; margin: 2.5cm; }
        body {
            font-family: Arial, "Liberation Sans", sans-serif;
            font-size: 11pt;
            line-height: 1.4;
            color: #000;
            margin: 0;
        }
        p                 { margin: 0 0 6pt 0; }
        h1                { font-size: 16pt; font-weight: bold; margin: 12pt 0 6pt; }
        h2                { font-size: 14pt; font-weight: bold; margin: 10pt 0 6pt; }
        h3                { font-size: 12pt; font-weight: bold; margin:  8pt 0 4pt; }
        h4, h5, h6        { font-size: 11pt; font-weight: bold; margin:  6pt 0 3pt; }
        table             { border-collapse: collapse; width: 100%; margin: 6pt 0; }
        td, th            { border: 1px solid #ccc; padding: 3pt 5pt; font-size: 10pt; }
        img               { max-width: 100%; height: auto; }
        ol, ul            { margin: 0 0 6pt; padding-left: 20pt; }
        li                { margin-bottom: 2pt; }
        .doc-header {
            font-size: 9pt; color: #555;
            border-bottom: 1px solid #ccc;
            padding-bottom: 4pt; margin-bottom: 12pt;
        }
        .doc-footer {
            font-size: 9pt; color: #555;
            border-top: 1px solid #ccc;
            padding-top: 4pt; margin-top: 12pt;
        }
    """

    html = (
        '<!DOCTYPE html><html lang="nl"><head>'
        '<meta charset="utf-8">'
        f'<style>{css}</style>'
        '</head><body>'
        f'{header_html}{body_html}{footer_html}'
        '</body></html>'
    )

    # ── HTML → PDF via WeasyPrint ─────────────────────────────────────────────
    weasyprint.HTML(string=html, base_url=output_dir).write_pdf(pdf_path)
    _validate_pdf(pdf_path, 'WeasyPrint: ')
    return pdf_path


def _convert_with_libreoffice(docx_path: str, output_dir: str) -> str:
    """
    Convert DOCX → PDF via LibreOffice headless subprocess.

    Highest layout fidelity but requires LibreOffice to be installed and a
    working VCL display plugin (SAL_USE_VCLPLUGIN=svp or xvfb on Linux).
    Used as a fallback when WeasyPrint is unavailable.
    """
    profile_dir = os.path.join(output_dir, f'lo_profile_{uuid.uuid4().hex}')
    os.makedirs(profile_dir, exist_ok=True)

    binary = _find_libreoffice()
    cmd    = _lo_cmd(binary, profile_dir, output_dir, docx_path)

    if sys.platform == 'darwin':
        result = subprocess.run(
            cmd, capture_output=True, text=True, timeout=120,
            env=_lo_env({'SAL_USE_VCLPLUGIN': 'svp'}),
        )
        if result.returncode != 0 and _is_display_error(result.stderr):
            result = subprocess.run(
                cmd, capture_output=True, text=True, timeout=120,
                env=_lo_env(),
            )
    else:
        env = _lo_env({'SAL_USE_VCLPLUGIN': 'svp'})
        env.pop('DISPLAY', None)
        env.pop('WAYLAND_DISPLAY', None)
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120, env=env)

        if result.returncode != 0 and _is_display_error(result.stderr):
            xvfb = shutil.which('xvfb-run')
            if xvfb:
                env2 = _lo_env()
                env2.pop('DISPLAY', None)
                env2.pop('WAYLAND_DISPLAY', None)
                result = subprocess.run(
                    [xvfb, '-a', '--server-args=-screen 0 1024x768x24'] + cmd,
                    capture_output=True, text=True, timeout=120, env=env2,
                )

    if result.returncode != 0:
        log.error(
            'LibreOffice conversion failed (rc=%d)\nSTDOUT: %s\nSTDERR: %s',
            result.returncode, result.stdout.strip(), result.stderr.strip(),
        )
        raise RuntimeError(
            f'LibreOffice-conversie mislukt (rc={result.returncode}):\n'
            f'{result.stderr.strip()}'
        )

    base     = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(output_dir, base + '.pdf')
    _validate_pdf(pdf_path, 'LibreOffice: ')
    return pdf_path


def convert_to_pdf(docx_path: str, output_dir: str) -> str:
    """
    Convert a .docx file to PDF.

    Strategy
    ────────
    1. WeasyPrint + mammoth  (primary)
       Pure Python — no external process, no display, no VCL plugin.
       Works headlessly on macOS (frozen .app) and Linux (Docker).

    2. LibreOffice headless  (fallback)
       Higher layout fidelity but requires LibreOffice to be installed
       with a working VCL plugin (SAL_USE_VCLPLUGIN=svp or xvfb).
       Only used when mammoth/weasyprint are not installed.
    """
    try:
        return _convert_with_weasyprint(docx_path, output_dir)
    except ImportError:
        log.info('mammoth/weasyprint niet beschikbaar, val terug op LibreOffice')
    except Exception as exc:
        log.warning('WeasyPrint-conversie mislukt (%s), val terug op LibreOffice', exc)

    return _convert_with_libreoffice(docx_path, output_dir)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def process(input_docx_path: str, output_pdf_path: str) -> None:
    """
    Full pipeline: load → clean → save cleaned docx → convert to PDF.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            doc = Document(input_docx_path)
        except Exception as exc:
            raise RuntimeError(
                'Kan het Word-document niet openen. '
                'Controleer of het bestand niet beschadigd of versleuteld is.'
            ) from exc

        remove_comments(doc)
        remove_highlighting(doc)
        remove_images_after_page_one(doc)

        cleaned_docx = os.path.join(tmpdir, 'cleaned.docx')
        doc.save(cleaned_docx)

        pdf_path = convert_to_pdf(cleaned_docx, tmpdir)
        shutil.move(pdf_path, output_pdf_path)
